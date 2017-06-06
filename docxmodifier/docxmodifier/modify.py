from docx import Document

def modify_paragraph(doc, replace_data):
    for p in doc.paragraphs:
        for old_text, new_text in replace_data.iteritems():
            if old_text in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    if old_text in inline[i].text:
                        text = inline[i].text.replace(old_text, new_text)
                        inline[i].text = text
                #print p.text
def replace_string(filename, data):
    replace_data = data["replace"]
    doc = Document(filename)
    modify_paragraph(doc, replace_data)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                modify_paragraph(cell, replace_data)
    doc.save('modify_input.docx')
    return 1
